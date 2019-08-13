# -*- coding: utf-8 -*-

import docx
import os
import comtypes.client
import pandas as pd

def add_lyrics_to_docx(doc, lyrics, song_title, album_title, artist, album_year,
                line_separator = "<br>"):
    """
    Creates a docx file that contains the given lyrics

    Args:
        doc: (Document) docx Document
        lyrics: (String)
        song_title: (String)
        album_title: (String)
        artist: (String)
        album_year: (String)
        lin_separator: (String) Used to split lyrics into separate lines
        
    Returns:
        - 
    
    Raises:
        - ReferenceError: When doc_path is not valid
    """
    
    doc.add_heading(song_title, 0)
    doc.add_heading(f"{artist} ({album_title} - {album_year})", 1)
    doc.add_paragraph("")

    for line in lyrics.split(line_separator):
        p = doc.add_paragraph(line.strip())
        p.paragraph_format.space_after = 0
    
    doc.add_page_break()
    
    return doc

def convert_to_pdf(doc_path, delete_docx=True):
    """
    Converts a given .docx file to PDF.

    Args:
        doc_path:    (String) Path to docx file, e.g. 'C:/Users/admin/test.docx'
        delete_docx: (Boolean) Determines whether the original docx file will
                     be deleted
    Returns:
        -
    
    Raises:
        ReferenceError: When doc_path is not valid
    """
    word = comtypes.client.CreateObject("Word.Application")
    
    if not os.path.isfile(doc_path):
        raise ReferenceError(f"Path not found: {str(doc_path)}")
        
    doc = word.Documents.Open(doc_path)
    doc.SaveAs(doc_path.replace("docx", "pdf"), FileFormat=17) # 17 = PDF
    doc.Close()
    word.Quit()
    if delete_docx:
        os.remove(doc_path)
    
if __name__ == "__main__":
    songs = pd.read_csv(os.getcwd() + "\\songs.csv")    
    songs[songs.Title.isna()]["Title"] = "45"
    
    
    PATH = os.getcwd() + "\\02_output\\test2.docx"
    doc = docx.Document()
    for idx, row in songs.iterrows():
        print(row["Title"])
        doc = add_lyrics_to_docx(doc, 
                                 lyrics=row["Lyrics"], 
                                 song_title=row["Title"],
                                 album_title=row["Album"],
                                 artist=row["Artist"],
                                 album_year=row["Year"],
                                 line_separator="\r\n")
    doc.save(PATH)
    #convert_to_pdf(PATH, delete_docx=True)